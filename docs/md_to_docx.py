from __future__ import annotations

from datetime import date
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_LOGO = ROOT / "favicon.png"
DEFAULT_THEME_COLOR = RGBColor(0x6D, 0x11, 0x16)
DEFAULT_SNAPSHOT_DATE = date.today().isoformat()


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    text = OxmlElement("w:t")
    text.text = " "
    run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_page_field(paragraph) -> None:
    add_field(paragraph, "PAGE")


def add_toc_field(paragraph) -> None:
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')


def clean_inline_markup(text: str) -> str:
    cleaned = re.sub(r"`([^`]+)`", r"\1", text)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    return cleaned.strip()


def flush_paragraph(doc: Document, lines: list[str]) -> None:
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        doc.add_paragraph(clean_inline_markup(text))


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [clean_inline_markup(cell.strip()) for cell in stripped.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return

    headers = rows[0]
    data_rows = [row for row in rows[2:] if any(cell.strip() for cell in row)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, cell in enumerate(headers):
        header_cells[idx].text = cell

    for row in data_rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(row):
            if idx < len(cells):
                cells[idx].text = cell


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)


def first_heading(markdown_text: str) -> str:
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return clean_inline_markup(stripped[2:])
    return "PensionApp Documentation"


def default_metadata(md_path: Path, markdown_text: str) -> dict:
    stem = md_path.stem.lower()
    audience = "General"
    description = "System reference"
    title = first_heading(markdown_text)

    if "technical" in stem or "system_documentation" in stem:
        audience = "Technical and implementation teams"
        description = "Architecture, data model, security, and operational reference"
    elif "user_admin" in stem or "manual" in stem:
        audience = "Operational users and administrators"
        description = "Feature usage guide and operational manual"

    return {
        "title": title,
        "subtitle": "UPS PensionsGo",
        "description": description,
        "audience": audience,
        "classification": "Internal Use",
        "version": "1.0",
        "status": "Draft for Review",
        "snapshot_date": DEFAULT_SNAPSHOT_DATE,
        "document_owner": "PensionsGo Project Team",
        "prepared_by": "OpenAI Codex",
        "logo_path": str(DEFAULT_LOGO),
        "theme_color": DEFAULT_THEME_COLOR,
        "revision_history": [
            {
                "version": "1.0",
                "date": DEFAULT_SNAPSHOT_DATE,
                "summary": "Initial branded documentation baseline generated from repository analysis.",
                "author": "OpenAI Codex",
                "approval": "Pending",
            }
        ],
        "approvals": [
            ("Prepared by", "Project Documentation / Engineering"),
            ("Reviewed by", "Technical Lead / Operations Lead"),
            ("Approved by", "Project Sponsor / Head of Pensions"),
        ],
    }


def add_cover_page(doc: Document, metadata: dict) -> None:
    logo_path = Path(str(metadata.get("logo_path") or "")).expanduser()
    color = metadata.get("theme_color", DEFAULT_THEME_COLOR)

    if logo_path.exists() and logo_path.stat().st_size > 0:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.3))

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(str(metadata.get("title", "PensionApp Documentation")))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = color

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(str(metadata.get("subtitle", "UPS PensionsGo")))
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(14)

    description_paragraph = doc.add_paragraph()
    description_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description_paragraph.add_run(str(metadata.get("description", "")))

    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    for label, value in [
        ("Audience", metadata.get("audience", "")),
        ("Document version", metadata.get("version", "")),
        ("Status", metadata.get("status", "")),
        ("Classification", metadata.get("classification", "")),
        ("Repository snapshot date", metadata.get("snapshot_date", "")),
        ("Document owner", metadata.get("document_owner", "")),
        ("Prepared by", metadata.get("prepared_by", "")),
    ]:
        row = info_table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)


def add_document_control_page(doc: Document, md_path: Path, metadata: dict) -> None:
    doc.add_page_break()
    doc.add_heading("Document Control", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"

    rows = [
        ("Document title", metadata.get("title", "")),
        ("Source markdown", md_path.name),
        ("Audience", metadata.get("audience", "")),
        ("Classification", metadata.get("classification", "")),
        ("Version", metadata.get("version", "")),
        ("Status", metadata.get("status", "")),
        ("Snapshot date", metadata.get("snapshot_date", "")),
        ("Prepared by", metadata.get("prepared_by", "")),
        ("Document owner", metadata.get("document_owner", "")),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = str(left)
        cells[1].text = str(right)


def add_revision_history_page(doc: Document, metadata: dict) -> None:
    doc.add_page_break()
    doc.add_heading("Revision History", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Version", "Date", "Summary", "Author", "Approval Status"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for item in metadata.get("revision_history", []):
        row = table.add_row().cells
        row[0].text = str(item.get("version", ""))
        row[1].text = str(item.get("date", ""))
        row[2].text = str(item.get("summary", ""))
        row[3].text = str(item.get("author", ""))
        row[4].text = str(item.get("approval", ""))


def add_approval_page(doc: Document, metadata: dict) -> None:
    doc.add_page_break()
    doc.add_heading("Approval and Sign-Off", level=1)
    doc.add_paragraph(
        "Use this page during review and handover to record document acceptance, ownership, and the latest approval state."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, header in enumerate(["Role", "Name / Team", "Signature", "Date"]):
        table.rows[0].cells[idx].text = header

    for role, default_name in metadata.get("approvals", []):
        row = table.add_row().cells
        row[0].text = str(role)
        row[1].text = str(default_name)
        row[2].text = ""
        row[3].text = ""


def add_toc_page(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    toc_note = doc.add_paragraph()
    toc_note.add_run("Note: ").bold = True
    toc_note.add_run("If the contents do not populate immediately, open the document in Word and update fields.")
    toc_paragraph = doc.add_paragraph()
    add_toc_field(toc_paragraph)


def render_markdown_body(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    para_buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith("```"):
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            code_paragraph = doc.add_paragraph()
            code_run = code_paragraph.add_run("\n".join(code_lines))
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(9)
            if i < len(lines) and lines[i].strip().startswith("```"):
                i += 1
            continue

        if stripped.startswith("#"):
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = clean_inline_markup(stripped[level:].strip())
            if level == 1 and heading_text == first_heading(markdown_text):
                i += 1
                continue
            doc.add_heading(heading_text, level=min(level, 3))
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                bullet = clean_inline_markup(lines[i].strip()[2:].strip())
                doc.add_paragraph(bullet, style="List Bullet")
                i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                doc.add_paragraph(clean_inline_markup(item), style="List Number")
                i += 1
            continue

        if stripped == "---":
            flush_paragraph(doc, para_buffer)
            para_buffer = []
            doc.add_page_break()
            i += 1
            continue

        para_buffer.append(stripped)
        i += 1

    flush_paragraph(doc, para_buffer)


def add_header_and_footer(doc: Document, metadata: dict) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.text = str(metadata.get("title", "PensionApp Documentation"))

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        add_page_field(footer)


def render_markdown_to_docx(md_path: Path, docx_path: Path, metadata: dict | None = None) -> None:
    markdown_text = md_path.read_text(encoding="utf-8", errors="ignore")
    meta = default_metadata(md_path, markdown_text)
    if metadata:
        meta.update(metadata)

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    ensure_styles(doc)
    add_cover_page(doc, meta)
    add_document_control_page(doc, md_path, meta)
    add_revision_history_page(doc, meta)
    add_approval_page(doc, meta)
    add_toc_page(doc)
    render_markdown_body(doc, markdown_text)
    add_header_and_footer(doc, meta)

    doc.core_properties.title = str(meta.get("title", "PensionApp Documentation"))
    doc.core_properties.subject = str(meta.get("description", "Documentation"))
    doc.core_properties.author = str(meta.get("prepared_by", "OpenAI Codex"))
    doc.core_properties.comments = f"Generated from {md_path.name} on {date.today().isoformat()}"
    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python md_to_docx.py <input.md> <output.docx>")
    render_markdown_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))

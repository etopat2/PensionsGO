from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Retirement_Type_Formula_Handout.docx"
LOGO = ROOT / "favicon.png"
BRAND = RGBColor(0x6D, 0x11, 0x16)
ACCENT = RGBColor(0x0E, 0x3A, 0x5D)
GOLD = RGBColor(0xC6, 0x9A, 0x2D)
LIGHT_GOLD = RGBColor(0xF7, 0xF0, 0xD7)
LIGHT_MAROON = RGBColor(0xF7, 0xEA, 0xEB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x20, 0x20, 0x20)


FORMULA_ROWS = [
    ("Mandatory Retirement", "mandatory", "Uses the Mandatory Retirement formula family"),
    ("Early Retirement", "early", "No benefits below 10 years; uses the Mandatory Retirement formula family at 10+ years"),
    ("Death", "death", "Gratuity is the higher of 3 x annual salary or the Mandatory Retirement gratuity; pension applies only at 10+ years"),
    ("Discharge (A.O.R)", "aor", "Same rule as Early Retirement"),
    ("Discharge (Medical)", "medical", "Uses the same rule as Death"),
    ("Discharge (Marriage)", "marriage", "Marriage gratuity only"),
    ("Discharge (C.B.E)", "cbe", "Short-service gratuity below 10 years; uses the Mandatory Retirement formula family at 10+ years"),
    ("Discharge (U.B.E)", "ube", "Short-service gratuity below 10 years; uses the Mandatory Retirement formula family at 10+ years"),
    ("Discharge (Public Interest)", "public", "Short-service gratuity below 10 years; uses the Mandatory Retirement formula family at 10+ years"),
    ("End of Contract", "contract", "Contract gratuity only"),
    ("Discharge (T.X)", "tx", "Uses the same rule as End of Contract"),
    ("Voluntary", "voluntary", "Uses the Mandatory Retirement formula family"),
    ("Old Age", "oldAge", "Uses the Mandatory Retirement formula family"),
    ("Abolition of Office", "abolition", "Uses the special 25% abolition formula"),
]

FORMULAS = [
    "Mandatory Retirement gratuity = (service x annual salary / 500) x (1/3) x 15",
    "Mandatory Retirement monthly pension = ((service x annual salary / 500) x (2/3)) / 12",
    "Mandatory Retirement full pension = (service x annual salary / 500) / 12",
    "Short-service gratuity = (service x annual salary x 10) / 500",
    "Discharge (Marriage) gratuity = (service x annual salary x 5) / 500",
    "End of Contract gratuity = 0.25 x annual salary x 2",
    "Abolition gratuity = ((service x annual salary / 500) x 0.25 x (1/3) x 15)",
    "Abolition monthly pension = ((service x annual salary / 500) x 0.25 x (2/3)) / 12",
    "Abolition full pension = ((service x annual salary / 500) x 0.25) / 12",
]

QA_CHECKS = [
    "Early Retirement and Discharge (A.O.R): same qualification route and same outputs for the same service, age profile, salary, and dates",
    "Discharge (C.B.E), Discharge (U.B.E), and Discharge (Public Interest) below 10 years: gratuity only",
    "Discharge (Marriage): gratuity only",
    "End of Contract and Discharge (T.X): same output",
    "Voluntary and Old Age: same output as Mandatory Retirement",
    "Abolition of Office: must not use the normal Mandatory Retirement pension values",
]

ALIASES = [
    "Contract Expired -> End of Contract (contract)",
    "Retirement by Death -> Death (death)",
    "At Own Request -> Discharge (A.O.R) (aor)",
    "Medical Grounds -> Discharge (Medical) (medical)",
    "Public Interest -> Discharge (Public Interest) (public)",
    "Discharge -> Discharge (C.B.E) (cbe)",
]


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 8.0, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_cell_margins(cell, top: int = 60, start: int = 90, bottom: int = 60, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_small_heading(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, "6D1116")
    set_table_cell_margins(cell, top=50, start=100, bottom=50, end=100)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = WHITE


def add_compact_bullets(doc: Document, items: list[str], *, size: float = 8.0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.14)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
        if not p.runs:
            run = p.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(size)
        else:
            p.runs[0].text = item


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8)

    banner = doc.add_table(rows=1, cols=2)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(8.9)
    banner.columns[1].width = Inches(1.2)
    title_cell = banner.rows[0].cells[0]
    logo_cell = banner.rows[0].cells[1]
    shade_cell(title_cell, "6D1116")
    shade_cell(logo_cell, "0E3A5D")
    set_table_cell_margins(title_cell, top=90, start=130, bottom=90, end=130)
    set_table_cell_margins(logo_cell, top=70, start=70, bottom=70, end=70)
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_after = Pt(0)
    t1 = tp.add_run("UGANDA PRISONS SERVICE\n")
    t1.bold = True
    t1.font.name = "Calibri"
    t1.font.size = Pt(8.5)
    t1.font.color.rgb = GOLD
    t2 = tp.add_run("Retirement Type Formula Handout")
    t2.bold = True
    t2.font.name = "Calibri"
    t2.font.size = Pt(19)
    t2.font.color.rgb = WHITE
    if LOGO.exists() and LOGO.stat().st_size > 0:
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(0.78))

    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    meta = [
        ("System", "UPS PensionsGo", "F7F0D7", BRAND),
        ("Audience", "Training & QA", "E8EFF5", ACCENT),
        ("Use", "Operational quick reference", "F7EAEB", BRAND),
    ]
    for idx, (label, value, fill, color) in enumerate(meta):
        cell = header_table.rows[0].cells[idx]
        shade_cell(cell, fill)
        set_table_cell_margins(cell, top=55, start=80, bottom=55, end=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(8)
        r1.font.color.rgb = color
        r2 = p.add_run(value)
        r2.font.name = "Calibri"
        r2.font.size = Pt(8)
        r2.font.color.rgb = INK

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(3)
    run = intro.add_run(
        "Use this sheet to confirm the approved retirement labels, benefit family applied, and the key exceptions staff should check before approval, registry entry, or QA sign-off."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8.2)

    add_small_heading(doc, "Canonical Retirement Types")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Display Label", "System Key", "Rule Summary"]
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], "0E3A5D")
        set_table_cell_margins(table.rows[0].cells[idx])
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8.2, color=WHITE)
    for row_index, (label, key, summary) in enumerate(FORMULA_ROWS, start=1):
        row = table.add_row().cells
        if row_index % 2 == 1:
            for cell in row:
                shade_cell(cell, "FBF7F1")
        set_cell_text(row[0], label, size=7.4)
        set_cell_text(row[1], key, size=7.4, color=BRAND)
        set_cell_text(row[2], summary, size=7.4)

    add_small_heading(doc, "Core Formulas")
    formulas_table = doc.add_table(rows=1, cols=2)
    formulas_table.style = "Table Grid"
    formulas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(["Formula / Rule", "QA Reminder"]):
        shade_cell(formulas_table.rows[0].cells[idx], "0E3A5D")
        set_table_cell_margins(formulas_table.rows[0].cells[idx])
        set_cell_text(formulas_table.rows[0].cells[idx], header, bold=True, size=8.2, color=WHITE)
    for idx in range(max(len(FORMULAS), len(QA_CHECKS))):
        row = formulas_table.add_row().cells
        if idx % 2 == 0:
            shade_cell(row[0], "F7EAEB")
            shade_cell(row[1], "F5F8FB")
        set_cell_text(row[0], FORMULAS[idx] if idx < len(FORMULAS) else "", size=7.25)
        set_cell_text(row[1], QA_CHECKS[idx] if idx < len(QA_CHECKS) else "", size=7.25)

    add_small_heading(doc, "Legacy Aliases Still Accepted")
    alias_table = doc.add_table(rows=1, cols=1)
    alias_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    alias_cell = alias_table.rows[0].cells[0]
    shade_cell(alias_cell, "F7F0D7")
    set_table_cell_margins(alias_cell, top=55, start=85, bottom=55, end=85)
    alias_line = alias_cell.paragraphs[0]
    alias_line.paragraph_format.space_before = Pt(0)
    alias_line.paragraph_format.space_after = Pt(0)
    alias_run = alias_line.add_run(" | ".join(ALIASES))
    alias_run.font.name = "Calibri"
    alias_run.font.size = Pt(7.8)
    alias_run.font.color.rgb = INK

    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    footer_table.columns[0].width = Inches(6.6)
    footer_table.columns[1].width = Inches(3.4)
    left_footer = footer_table.rows[0].cells[0]
    right_footer = footer_table.rows[0].cells[1]
    shade_cell(left_footer, "6D1116")
    shade_cell(right_footer, "0E3A5D")
    set_table_cell_margins(left_footer, top=45, start=85, bottom=45, end=85)
    set_table_cell_margins(right_footer, top=45, start=85, bottom=45, end=85)
    lp = left_footer.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run("Reference baseline: shared calculator + workflow + registry benefit engine")
    lr.font.name = "Calibri"
    lr.font.size = Pt(7.2)
    lr.font.color.rgb = WHITE
    rp = right_footer.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run("Internal training handout")
    rr.bold = True
    rr.font.name = "Calibri"
    rr.font.size = Pt(7.2)
    rr.font.color.rgb = GOLD

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
